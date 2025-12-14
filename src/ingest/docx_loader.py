from docx import Document
from .base_extractor import BaseExtractor

class DOCXExtractor(BaseExtractor):
    def __init__(self):
        super().__init__()
        
    def extract(self, file_path: str) -> str:
        file_exists = super().validate(file_path)
        if not file_exists:
            raise FileNotFoundError(f"File not found {file_path}")
        
        extracted_text = ""
        doc = Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip() != "":
                paragraphs.append(para.text.strip())

        extracted_text = "\n".join(paragraphs)
        
                
        for tab in doc.tables:
            table_data = []
            for row in tab.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                extracted_row_text = "|".join(row_data)
                table_data.append(extracted_row_text)
            extracted_text += "\n".join(table_data)
        
        return extracted_text
            
if __name__ == "__main__":
    extractor = DOCXExtractor()
    path = "data/raw/sample.docx"
    result = extractor.extract(path)
    print(result[:500])
    
                